from typing import List, Tuple, Optional
from dataclasses import dataclass
import os


@dataclass
class ParsedPage:
    page_number: int
    text: str
    headings: Optional[List[Tuple[int, str]]]


@dataclass
class ParsedDocument:
    pages: List[ParsedPage]
    full_text: str
    headings_hierarchy: List[Tuple[int, str, int]]


def parse_pdf(file_path: str) -> ParsedDocument:
    import pdfplumber

    pages: List[ParsedPage] = []
    all_headings = []
    full_text_parts = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            text = _clean_pdf_text(text)
            headings = _extract_pdf_headings(text, page_num)
            pages.append(ParsedPage(
                page_number=page_num,
                text=text,
                headings=headings
            ))
            full_text_parts.append(text)
            for level, heading_text in headings:
                all_headings.append((level, heading_text, page_num))

    full_text = "\n\n".join(full_text_parts)
    return ParsedDocument(
        pages=pages,
        full_text=full_text,
        headings_hierarchy=all_headings
    )


def _clean_pdf_text(text: str) -> str:
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if len(stripped) < 10 and (
            stripped.isdigit() or
            stripped.lower().startswith("page ") or
            stripped.lower().startswith("第") and stripped.endswith("页")
        ):
            continue
        if len(stripped) < 30 and (
            "copyright" in stripped.lower() or
            "confidential" in stripped.lower()
        ):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def _extract_pdf_headings(text: str, page_num: int) -> List[Tuple[int, str]]:
    headings = []
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.isupper() and len(stripped.split()) <= 10 and len(stripped) > 2:
            headings.append((2, stripped))
            continue
        import re
        h1_match = re.match(r"^\d+\.\s+[A-Z\u4e00-\u9fff].*", stripped)
        if h1_match and len(stripped) < 80:
            headings.append((1, stripped))
            continue
        h2_match = re.match(r"^\d+\.\d+\s+", stripped)
        if h2_match and len(stripped) < 80:
            headings.append((2, stripped))
            continue
        h3_match = re.match(r"^\d+\.\d+\.\d+\s+", stripped)
        if h3_match and len(stripped) < 80:
            headings.append((3, stripped))
            continue
    return headings


def parse_docx(file_path: str) -> ParsedDocument:
    from docx import Document

    pages: List[ParsedPage] = []
    all_headings = []
    full_text_parts = []
    page_num = 1
    current_page_text = []
    page_headings = []

    doc = Document(file_path)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        if not text:
            continue

        level = 0
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
        elif style_name == "Title":
            level = 0

        if level > 0:
            current_page_text.append(text)
            page_headings.append((level, text))
            all_headings.append((level, text, page_num))
        else:
            current_page_text.append(text)

        if len("\n".join(current_page_text)) > 3000:
            pages.append(ParsedPage(
                page_number=page_num,
                text="\n".join(current_page_text),
                headings=page_headings
            ))
            full_text_parts.append("\n".join(current_page_text))
            page_num += 1
            current_page_text = []
            page_headings = []

    if current_page_text:
        pages.append(ParsedPage(
            page_number=page_num,
            text="\n".join(current_page_text),
            headings=page_headings
        ))
        full_text_parts.append("\n".join(current_page_text))

    full_text = "\n\n".join(full_text_parts)
    return ParsedDocument(
        pages=pages,
        full_text=full_text,
        headings_hierarchy=all_headings
    )


def parse_txt(file_path: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    pages: List[ParsedPage] = []
    all_headings = []

    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

    chars_per_page = 3000
    current_page_text = []
    page_headings = []
    page_num = 1
    char_count = 0

    for para in paragraphs:
        lines = para.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if len(stripped) < 80 and (
                stripped.isupper() or
                stripped.startswith("#") or
                stripped.startswith("第") and "章" in stripped
            ):
                level = 1 if len(stripped.split()) <= 10 else 2
                page_headings.append((level, stripped))
                all_headings.append((level, stripped, page_num))

            current_page_text.append(line)
            char_count += len(line)

        if char_count > chars_per_page:
            pages.append(ParsedPage(
                page_number=page_num,
                text="\n".join(current_page_text),
                headings=page_headings
            ))
            page_num += 1
            current_page_text = []
            page_headings = []
            char_count = 0

    if current_page_text:
        pages.append(ParsedPage(
            page_number=page_num,
            text="\n".join(current_page_text),
            headings=page_headings
        ))

    return ParsedDocument(
        pages=pages,
        full_text=content,
        headings_hierarchy=all_headings
    )


def parse_markdown(file_path: str) -> ParsedDocument:
    import markdown
    from bs4 import BeautifulSoup

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        md_content = f.read()

    html = markdown.markdown(md_content, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    pages: List[ParsedPage] = []
    all_headings = []

    full_text_parts = []
    current_page_text = []
    page_headings = []
    page_num = 1

    for element in soup.children:
        text = ""
        level = 0

        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            text = element.get_text().strip()
            page_headings.append((level, text))
            all_headings.append((level, text, page_num))
        elif element.name:
            text = element.get_text().strip()

        if text:
            current_page_text.append(text)

        if len("\n".join(current_page_text)) > 3000:
            pages_text = "\n".join(current_page_text)
            pages.append(ParsedPage(
                page_number=page_num,
                text=pages_text,
                headings=page_headings
            ))
            full_text_parts.append(pages_text)
            page_num += 1
            current_page_text = []
            page_headings = []

    if current_page_text:
        pages_text = "\n".join(current_page_text)
        pages.append(ParsedPage(
            page_number=page_num,
            text=pages_text,
            headings=page_headings
        ))
        full_text_parts.append(pages_text)

    full_text = "\n\n".join(full_text_parts)
    return ParsedDocument(
        pages=pages,
        full_text=full_text,
        headings_hierarchy=all_headings
    )


def parse_document(file_path: str) -> ParsedDocument:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx"]:
        return parse_docx(file_path)
    elif ext in [".txt"]:
        return parse_txt(file_path)
    elif ext in [".md", ".markdown"]:
        return parse_markdown(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
